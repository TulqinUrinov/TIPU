import io
import os
import tempfile
import subprocess
import qrcode
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
from data.student.models import Student
from data.contract.models import InstallmentPayment
from data.file.models import ContractFiles, Files


def generate_contract(student: Student) -> ContractFiles:
    """
    Student uchun shartnoma yaratadi va DB ga saqlaydi.
    PDF faylni ContractFiles modeliga yozib beradi.
    """

    # Avval DB da mavjud bo'lganini tekshiramiz
    existing_contract = ContractFiles.objects.filter(student=student).first()
    if existing_contract and existing_contract.file and os.path.exists(existing_contract.file.path):
        return existing_contract

    # Student turini aniqlash
    template_type = "MUQOBIL" if hasattr(student, "user_account") else "HEMIS"

    try:
        file = Files.objects.get(file_type=template_type)
    except Files.DoesNotExist:
        raise Exception(f"{template_type} shablon topilmadi")

    contract = student.contract.first()

    # O'rnini bosadigan matnlar
    replacements = {
        "{filial}": "Bosh filial",
        "{name}": str(student.full_name),
        "{mode}": str(student.education_form),
        "{delta}": str(5 if student.education_form == "Sirtqi" else 4),
        "{course}": str(student.course),
        "{faculty}": str(student.specialization.name),
        "{price}": str(contract.period_amount_dt) if contract else "",
        "{jshir}": str(student.jshshir),
        "{phone}": str(student.user_account.phone_number) if hasattr(student, "user_account") else str(
            student.phone_number or ""),
    }

    # QR kod
    qr_url = f"{settings.SITE_URL}/contracts/{student.id}/download/"
    qr_img = qrcode.make(qr_url)
    qr_stream = io.BytesIO()
    qr_img.save(qr_stream, format="PNG")
    qr_stream.seek(0)

    # DOCX yaratish
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        doc = Document(file.file.path)

        # Paragraph va table ichidagi matnlarni almashtirish
        for paragraph in doc.paragraphs:
            replace_text_preserving_format(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_preserving_format(paragraph, replacements)

        if template_type == "MUQOBIL":
            fill_installment_table(doc, student)

        # QR kod joylash
        for paragraph in doc.paragraphs:
            if "{qr}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{qr}", "")
                run = paragraph.add_run()
                run.add_picture(qr_stream, width=Inches(1.5))

        doc.save(tmp_docx.name)
        output_docx = tmp_docx.name

    # PDF ga convert qilish
    subprocess.run([
        "libreoffice", "--headless",
        "--convert-to", "pdf",
        "--outdir", tempfile.gettempdir(),
        output_docx
    ], check=True)

    pdf_path = output_docx.replace(".docx", ".pdf")

    # DB ga yozish
    with open(pdf_path, "rb") as f:
        contract_file = ContractFiles.objects.create(student=student)
        contract_file.file.save(f"contract_{student.id}.pdf", f)

    # Vaqtinchalik fayllarni tozalash
    os.remove(output_docx)
    if os.path.exists(pdf_path):
        os.remove(pdf_path)

    return contract_file


def replace_text_preserving_format(paragraph, replacements):
    for run in paragraph.runs:
        original_text = run.text
        new_text = original_text
        for key, value in replacements.items():
            if key in new_text:
                new_text = new_text.replace(key, value)
        if new_text != original_text:
            run.text = new_text


def fill_installment_table(doc, student):
    try:
        installment_payment = InstallmentPayment.objects.filter(student=student).first()
        if not installment_payment:
            return

        installment_data = installment_payment.installment_payments

        for table in doc.tables:
            if len(table.columns) >= 4:
                header_row = table.rows[0]
                set_cell_text_preserve_format(header_row.cells[0], "To'lanishi kerak")
                set_cell_text_preserve_format(header_row.cells[1], "Muddat")
                set_cell_text_preserve_format(header_row.cells[2], "To'langan summa")
                set_cell_text_preserve_format(header_row.cells[3], "Qolgan summa")

                while len(table.rows) > 5:
                    table._tbl.remove(table.rows[-1]._tr)

                while len(table.rows) < 5:
                    table.add_row()

                for i in range(4):
                    row = table.rows[i + 1]
                    if i < len(installment_data):
                        installment = installment_data[i]
                        set_cell_text_preserve_format(row.cells[0], f"{float(installment.get('amount', 0)):,.0f} so'm")
                        payment_date = installment.get('payment_date', '')
                        if payment_date:
                            try:
                                date_obj = datetime.strptime(payment_date, '%Y-%m-%d')
                                formatted_date = date_obj.strftime('%d.%m.%Y')
                                set_cell_text_preserve_format(row.cells[1], formatted_date)
                            except:
                                set_cell_text_preserve_format(row.cells[1], payment_date)
                        paid = float(installment.get('amount', 0)) - float(installment.get('left', 0))
                        set_cell_text_preserve_format(row.cells[2], f"{paid:,.0f} so'm")
                        set_cell_text_preserve_format(row.cells[3], f"{float(installment.get('left', 0)):,.0f} so'm")
                    else:
                        for j in range(4):
                            set_cell_text_preserve_format(row.cells[j], "")
                break
    except Exception as e:
        print(f"To'lov jadvalini to'ldirishda xato: {e}")


def set_cell_text_preserve_format(cell, text):
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    new_paragraph = cell.add_paragraph()
    run = new_paragraph.add_run(text)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
