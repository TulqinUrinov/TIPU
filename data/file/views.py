import io
import os
import subprocess
import tempfile
from datetime import datetime

from django.conf import settings
from django.http import FileResponse, Http404
from docx.enum.text import WD_ALIGN_PARAGRAPH

from rest_framework import generics, views, viewsets
from rest_framework.generics import ListAPIView

from docx import Document
from docx.shared import Inches

import qrcode

from data.common.permission import IsAuthenticatedUserType
from data.file.models import Files, ContractFiles
from data.file.serializers import FileSerializer, FileUploadSerializer
from data.payment.models import InstallmentPayment
from data.student.models import Student


class ImportHistoryAPIView(ListAPIView):
    permission_classes = [IsAuthenticatedUserType]
    queryset = Files.objects.all().order_by("-created_at")
    serializer_class = FileSerializer


# Barcha Filelar
class FileViewSet(viewsets.ModelViewSet):
    queryset = Files.objects.all()
    serializer_class = FileUploadSerializer
    permission_classes = [IsAuthenticatedUserType]


# faqat shablon
class SpecialDocsListApiView(generics.ListAPIView):
    serializer_class = FileUploadSerializer
    permission_classes = [IsAuthenticatedUserType]

    def get_queryset(self):
        return Files.objects.filter(file_type__in=["MUQOBIL", "HEMIS"])


class ContractDownloadApiView(views.APIView):
    """Shartnomani PDF shaklida yuklab beruvchi API"""
    permission_classes = [IsAuthenticatedUserType]

    def get(self, request, pk):
        try:
            student = Student.objects.get(pk=pk)
        except Student.DoesNotExist:
            raise Http404("Student not found")

        # Avval DB dan mavjud shartnomani topamiz
        existing_contract = ContractFiles.objects.filter(student=student).first()

        # Agar mavjud bo'lsa va fayl mavjud bo'lsa → qaytarish
        if existing_contract and existing_contract.file and os.path.exists(existing_contract.file.path):
            return FileResponse(
                existing_contract.file.open("rb"),
                as_attachment=True,
                filename=os.path.basename(existing_contract.file.name)
            )

        # Student turini aniqlash
        template_type = "MUQOBIL" if hasattr(student, "user_account") else "HEMIS"

        # Shablonni olish
        try:
            file = Files.objects.get(file_type=template_type)
        except Files.DoesNotExist:
            raise Http404(f"{template_type} shablon topilmadi")

        contract = student.contract.first()

        # To'ldirish uchun ma'lumotlar
        replacements = {
            "{filial}": "Bosh filial",  # Kerak bo'lsa, student ma'lumotlaridan oling
            "{name}": str(student.full_name),
            "{mode}": str(student.education_form),
            "{delta}": str(5 if student.education_form == "Sirtqi" else 4),
            "{course}": str(student.course),
            "{faculty}": str(student.specialization.name),
            "{price}": str(contract.period_amount_dt) if contract else "",
            "{jshir}": str(student.jshshir),
            "{phone}": str(student.user_account.phone_number) if hasattr(student, "user_account") else str(
                student.phone_number or ""),
        }

        # QR kod yaratish
        qr_url = f"{settings.SITE_URL}/contracts/{student.id}/download/"
        qr_img = qrcode.make(qr_url)
        qr_stream = io.BytesIO()
        qr_img.save(qr_stream, format="PNG")
        qr_stream.seek(0)

        # Vaqtinchalik docx yaratish
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            doc = Document(file.file.path)

            # Paragraphlardagi kalit so'zlarni almashtirish (formatni saqlab qolgan holda)
            for paragraph in doc.paragraphs:
                self.replace_text_preserving_format(paragraph, replacements)

            # Jadvaldagi kalit so'zlarni almashtirish (formatni saqlab qolgan holda)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_preserving_format(paragraph, replacements)

            # MUQOBIL shartnomasi bo'lsa, to'lov jadvalini to'ldirish
            if template_type == "MUQOBIL":
                self.fill_installment_table(doc, student)

            # QR kod qo'shish
            for paragraph in doc.paragraphs:
                if "{qr}" in paragraph.text:
                    # {qr} ni olib tashlash
                    paragraph.text = paragraph.text.replace("{qr}", "")
                    # QR kodni qo'shish
                    run = paragraph.add_run()
                    run.add_picture(qr_stream, width=Inches(1.5))

            doc.save(tmp_docx.name)
            output_docx = tmp_docx.name

        # LibreOffice orqali PDF ga o'tkazish
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tempfile.gettempdir(),
            output_docx
        ], check=True)

        pdf_path = output_docx.replace(".docx", ".pdf")

        # DB ga PDF saqlash
        with open(pdf_path, "rb") as f:
            contract_file = ContractFiles.objects.create(student=student)
            contract_file.file.save(f"contract_{student.id}.pdf", f)

        # Vaqtinchalik fayllarni o'chirish
        os.remove(output_docx)
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

        return FileResponse(
            contract_file.file.open("rb"),
            as_attachment=True,
            filename=f"contract_{student.id}.pdf"
        )

    def replace_text_preserving_format(self, paragraph, replacements):
        """Formatlashni saqlab qolgan holda matnni almashtirish"""
        # Paragraphdagi barcha runlarni tekshiramiz
        for run in paragraph.runs:
            original_text = run.text
            new_text = original_text

            # Har bir kalit so'zni almashtiramiz
            for key, value in replacements.items():
                if key in new_text:
                    new_text = new_text.replace(key, value)

            # Agar matn o'zgarganga, runni yangilaymiz
            if new_text != original_text:
                run.text = new_text

    def fill_installment_table(self, doc, student):
        """To'lov jadvalini to'ldirish"""
        try:
            # Studentning to'lov ma'lumotlarini olish
            installment_payment = InstallmentPayment.objects.filter(student=student).first()

            if not installment_payment:
                return

            installment_data = installment_payment.installment_payments

            # Jadvalni topish (4 ustunli jadval)
            for table in doc.tables:
                if len(table.columns) >= 4:  # Kamida 4 ustunli jadval
                    # Jadvaldagi mavjud qatorlarni tekshiramiz
                    # Sarlavha qatorini topamiz
                    header_found = False
                    for i, row in enumerate(table.rows):
                        row_text = " ".join([cell.text for cell in row.cells])
                        if "To'lanishi kerak" in row_text or "Muddat" in row_text or "Tolangan summa" in row_text or "Qolgan summa" in row_text:
                            header_found = True
                            break

                    # Agar sarlavha qatori topilmasa, birinchi qatorni sarlavha deb hisoblaymiz
                    if not header_found:
                        header_row = table.rows[0]
                        # Sarlavhalarni to'g'ri joylashtiramiz
                        if len(header_row.cells) >= 4:
                            self.set_cell_text_preserve_format(header_row.cells[0], "To'lanishi kerak")
                            self.set_cell_text_preserve_format(header_row.cells[1], "Muddat")
                            self.set_cell_text_preserve_format(header_row.cells[2], "To'langan summa")
                            self.set_cell_text_preserve_format(header_row.cells[3], "Qolgan summa")

                    # Faqat 4 ta to'lov qatori bo'lishini ta'minlaymiz (sarlavhadan keyin)
                    while len(table.rows) > 5:  # 1 sarlavha + 4 to'lov = 5 qator
                        table._tbl.remove(table.rows[-1]._tr)

                    # Yetarli qator yo'q bo'lsa, qo'shamiz
                    while len(table.rows) < 5:
                        table.add_row()

                    # Har bir to'lov ma'lumotini qatorlarga joylashtirish
                    for i in range(4):  # 4 ta to'lov qatori
                        if i < len(installment_data):
                            installment = installment_data[i]
                            row = table.rows[i + 1]  # 0-index sarlavha, shuning uchun i+1

                            # To'lov miqdori
                            if len(row.cells) > 0:
                                self.set_cell_text_preserve_format(row.cells[0],
                                                                   f"{float(installment.get('amount', 0)):,.0f} so'm")

                            # Muddat
                            if len(row.cells) > 1:
                                payment_date = installment.get('payment_date', '')
                                if payment_date:
                                    try:
                                        date_obj = datetime.strptime(payment_date, '%Y-%m-%d')
                                        formatted_date = date_obj.strftime('%d.%m.%Y')
                                        self.set_cell_text_preserve_format(row.cells[1], formatted_date)
                                    except:
                                        self.set_cell_text_preserve_format(row.cells[1], payment_date)

                            # To'langan summa
                            if len(row.cells) > 2:
                                paid = float(installment.get('amount', 0)) - float(installment.get('left', 0))
                                self.set_cell_text_preserve_format(row.cells[2], f"{paid:,.0f} so'm")

                            # Qolgan summa
                            if len(row.cells) > 3:
                                self.set_cell_text_preserve_format(row.cells[3],
                                                                   f"{float(installment.get('left', 0)):,.0f} so'm")
                        else:
                            # Agar ma'lumot yetarli bo'lmasa, bo'sh qoldirish
                            row = table.rows[i + 1]
                            for j in range(4):
                                if j < len(row.cells):
                                    self.set_cell_text_preserve_format(row.cells[j], "")

                    break

        except Exception as e:
            print(f"To'lov jadvalini to'ldirishda xato: {e}")

    def set_cell_text_preserve_format(self, cell, text):
        """Hujayra matnini o'rnatish (formatni saqlab qolgan holda)"""
        # Hujayradagi barcha paragraphlarni o'chirib, yangisini qo'shamiz
        for paragraph in cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)

        # Yangi paragraph qo'shamiz
        new_paragraph = cell.add_paragraph()

        # Formatni saqlash uchun yangi run yaratamiz
        run = new_paragraph.add_run(text)

        # Asl formatni saqlash (agar mavjud bo'lsa)
        if cell.paragraphs and cell.paragraphs[0].runs:
            original_run = cell.paragraphs[0].runs[0]
            run.bold = original_run.bold
            run.italic = original_run.italic
            run.font.size = original_run.font.size
            run.font.name = original_run.font.name

        # Markazga tekislash
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# class ContractDownloadApiView(views.APIView):
#     """Shartnomani PDF shaklida yuklab beruvchi API"""
#     permission_classes = [IsAuthenticatedUserType]
#
#     def get(self, request, pk):
#         try:
#             student = Student.objects.get(pk=pk)
#         except Student.DoesNotExist:
#             raise Http404("Student not found")
#
#         # Avval DB dan mavjud shartnomani topamiz
#         existing_contract = ContractFiles.objects.filter(student=student).first()
#
#         # Agar mavjud bo'lsa va fayl mavjud bo'lsa → qaytarish
#         if existing_contract and existing_contract.file and os.path.exists(existing_contract.file.path):
#             return FileResponse(
#                 existing_contract.file.open("rb"),
#                 as_attachment=True,
#                 filename=os.path.basename(existing_contract.file.name)
#             )
#
#         # Student turini aniqlash
#         template_type = "MUQOBIL" if hasattr(student, "user_account") else "HEMIS"
#
#         # Shablonni olish
#         try:
#             file = Files.objects.get(file_type=template_type)
#         except Files.DoesNotExist:
#             raise Http404(f"{template_type} shablon topilmadi")
#
#         contract = student.contract.first()
#
#         # To'ldirish uchun ma'lumotlar
#         replacements = {
#             "{filial}": "Bosh filial",  # Kerak bo'lsa, student ma'lumotlaridan oling
#             "{name}": str(student.full_name),
#             "{mode}": str(student.education_form),
#             "{delta}": str(5 if student.education_form == "Sirtqi" else 4),
#             "{course}": str(student.course),
#             "{faculty}": str(student.specialization.name),
#             "{price}": str(contract.period_amount_dt) if contract else "",
#             "{jshir}": str(student.jshshir),
#             "{phone}": str(student.user_account.phone_number) if hasattr(student, "user_account") else str(
#                 student.phone_number or ""),
#         }
#
#         # QR kod yaratish
#         qr_url = f"{settings.SITE_URL}/contracts/{student.id}/download/"
#         qr_img = qrcode.make(qr_url)
#         qr_stream = io.BytesIO()
#         qr_img.save(qr_stream, format="PNG")
#         qr_stream.seek(0)
#
#         # Vaqtinchalik docx yaratish
#         with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
#             doc = Document(file.file.path)
#
#             # Paragraphlardagi kalit so'zlarni almashtirish (formatni saqlab qolgan holda)
#             for paragraph in doc.paragraphs:
#                 self.replace_text_preserving_format(paragraph, replacements)
#
#             # Jadvaldagi kalit so'zlarni almashtirish (formatni saqlab qolgan holda)
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         for paragraph in cell.paragraphs:
#                             self.replace_text_preserving_format(paragraph, replacements)
#
#             # MUQOBIL shartnomasi bo'lsa, to'lov jadvalini to'ldirish
#             if template_type == "MUQOBIL":
#                 self.fill_installment_table(doc, student)
#
#             # QR kod qo'shish
#             for paragraph in doc.paragraphs:
#                 if "{qr}" in paragraph.text:
#                     # {qr} ni olib tashlash
#                     paragraph.text = paragraph.text.replace("{qr}", "")
#                     # QR kodni qo'shish
#                     run = paragraph.add_run()
#                     run.add_picture(qr_stream, width=Inches(1.5))
#
#             doc.save(tmp_docx.name)
#             output_docx = tmp_docx.name
#
#         # LibreOffice orqali PDF ga o'tkazish
#         subprocess.run([
#             "libreoffice",
#             "--headless",
#             "--convert-to", "pdf",
#             "--outdir", tempfile.gettempdir(),
#             output_docx
#         ], check=True)
#
#         pdf_path = output_docx.replace(".docx", ".pdf")
#
#         # DB ga PDF saqlash
#         with open(pdf_path, "rb") as f:
#             contract_file = ContractFiles.objects.create(student=student)
#             contract_file.file.save(f"contract_{student.id}.pdf", f)
#
#         # Vaqtinchalik fayllarni o'chirish
#         os.remove(output_docx)
#         if os.path.exists(pdf_path):
#             os.remove(pdf_path)
#
#         return FileResponse(
#             contract_file.file.open("rb"),
#             as_attachment=True,
#             filename=f"contract_{student.id}.pdf"
#         )
#
#     def replace_text_preserving_format(self, paragraph, replacements):
#         """Formatlashni saqlab qolgan holda matnni almashtirish"""
#         # Paragraphdagi barcha runlarni tekshiramiz
#         for run in paragraph.runs:
#             original_text = run.text
#             new_text = original_text
#
#             # Har bir kalit so'zni almashtiramiz
#             for key, value in replacements.items():
#                 if key in new_text:
#                     new_text = new_text.replace(key, value)
#
#             # Agar matn o'zgarganga, runni yangilaymiz
#             if new_text != original_text:
#                 run.text = new_text
#
#     def fill_installment_table(self, doc, student):
#         """To'lov jadvalini to'ldirish"""
#         try:
#             # Studentning to'lov ma'lumotlarini olish
#             installment_payment = InstallmentPayment.objects.filter(student=student).first()
#
#             if not installment_payment:
#                 return
#
#             installment_data = installment_payment.installment_payments
#
#             # Jadvalni topish (4 ustunli jadval)
#             for table in doc.tables:
#                 if len(table.columns) >= 4:  # Kamida 4 ustunli jadval
#                     # Jadvaldagi mavjud qatorlarni saqlab qolamiz (sarlavha qatorini)
#                     # Yangi qatorlarni qo'shamiz
#                     for i, installment in enumerate(installment_data):
#                         # Jadvalga yangi qator qo'shamiz
#                         if i >= len(table.rows) - 1:  # Agar jadvalda yetarlicha qator yo'q bo'lsa
#                             table.add_row()
#
#                         # Qatorni to'ldirish
#                         row = table.rows[i + 1]  # 0-index sarlavha, shuning uchun i+1
#
#                         # To'lov miqdori
#                         if len(row.cells) > 0:
#                             self.set_cell_text(row.cells[0], f"{float(installment.get('amount', 0)):,.0f} so'm")
#
#                         # Muddat
#                         if len(row.cells) > 1:
#                             payment_date = installment.get('payment_date', '')
#                             if payment_date:
#                                 # Sanani formatlash
#                                 try:
#                                     date_obj = datetime.strptime(payment_date, '%Y-%m-%d')
#                                     formatted_date = date_obj.strftime('%d.%m.%Y')
#                                     self.set_cell_text(row.cells[1], formatted_date)
#                                 except:
#                                     self.set_cell_text(row.cells[1], payment_date)
#
#                         # To'flangan summa
#                         if len(row.cells) > 2:
#                             paid = float(installment.get('amount', 0)) - float(installment.get('left', 0))
#                             self.set_cell_text(row.cells[2], f"{paid:,.0f} so'm")
#
#                         # Qolgan summa
#                         if len(row.cells) > 3:
#                             self.set_cell_text(row.cells[3], f"{float(installment.get('left', 0)):,.0f} so'm")
#
#                     break
#
#         except Exception as e:
#             print(f"To'lov jadvalini to'ldirishda xato: {e}")
#
#     def set_cell_text(self, cell, text):
#         """Hujayra matnini o'rnatish"""
#         # Hujayradagi barcha paragraphlarni o'chirib, yangisini qo'shamiz
#         for paragraph in cell.paragraphs:
#             p = paragraph._element
#             p.getparent().remove(p)
#
#         # Yangi paragraph qo'shamiz
#         new_paragraph = cell.add_paragraph()
#         new_paragraph.text = text
#
#         # Markazga tekislash
#         for paragraph in cell.paragraphs:
#             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# class ContractDownloadApiView(views.APIView):
#     """Shartnomani PDF shaklida yuklab beruvchi API"""
#     permission_classes = [IsAuthenticatedUserType]
#
#     def get(self, request, pk):
#         try:
#             student = Student.objects.get(pk=pk)
#         except Student.DoesNotExist:
#             raise Http404("Student not found")
#
#         # Avval DB dan mavjud shartnomani topamiz
#         existing_contract = ContractFiles.objects.filter(student=student).first()
#
#         # Agar mavjud bo'lsa va fayl mavjud bo'lsa → qaytarish
#         if existing_contract and existing_contract.file and os.path.exists(existing_contract.file.path):
#             return FileResponse(
#                 existing_contract.file.open("rb"),
#                 as_attachment=True,
#                 filename=os.path.basename(existing_contract.file.name)
#             )
#
#         # Student turini aniqlash
#         template_type = "MUQOBIL" if hasattr(student, "user_account") else "HEMIS"
#
#         # Shablonni olish
#         try:
#             file = Files.objects.get(file_type=template_type)
#         except Files.DoesNotExist:
#             raise Http404(f"{template_type} shablon topilmadi")
#
#         contract = student.contract.first()
#
#         # To'ldirish uchun ma'lumotlar
#         replacements = {
#             "{filial}": "Bosh filial",  # Kerak bo'lsa, student ma'lumotlaridan oling
#             "{name}": str(student.full_name),
#             "{mode}": str(student.education_form),
#             "{delta}": str(5 if student.education_form == "Sirtqi" else 4),
#             "{course}": str(student.course),
#             "{faculty}": str(student.specialization.name),
#             "{price}": str(contract.period_amount_dt) if contract else "",
#             "{jshir}": str(student.jshshir),
#             "{phone}": str(student.user_account.phone_number) if hasattr(student, "user_account") else str(
#                 student.phone_number or ""),
#         }
#
#         # QR kod yaratish
#         qr_url = f"{settings.SITE_URL}/contracts/{student.id}/download/"
#         qr_img = qrcode.make(qr_url)
#         qr_stream = io.BytesIO()
#         qr_img.save(qr_stream, format="PNG")
#         qr_stream.seek(0)
#
#         # Vaqtinchalik docx yaratish
#         with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
#             doc = Document(file.file.path)
#
#             # Paragraphlardagi kalit so'zlarni almashtirish (formatni saqlab qolgan holda)
#             for paragraph in doc.paragraphs:
#                 self.replace_text_preserving_format(paragraph, replacements)
#
#             # Jadvaldagi kalit so'zlarni almashtirish (formatni saqlab qolgan holda)
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         for paragraph in cell.paragraphs:
#                             self.replace_text_preserving_format(paragraph, replacements)
#
#             # QR kod qo'shish
#             for paragraph in doc.paragraphs:
#                 if "{qr}" in paragraph.text:
#                     # {qr} ni olib tashlash
#                     paragraph.text = paragraph.text.replace("{qr}", "")
#                     # QR kodni qo'shish
#                     run = paragraph.add_run()
#                     run.add_picture(qr_stream, width=Inches(1.5))
#
#             doc.save(tmp_docx.name)
#             output_docx = tmp_docx.name
#
#         # LibreOffice orqali PDF ga o'tkazish
#         subprocess.run([
#             "libreoffice",
#             "--headless",
#             "--convert-to", "pdf",
#             "--outdir", tempfile.gettempdir(),
#             output_docx
#         ], check=True)
#
#         pdf_path = output_docx.replace(".docx", ".pdf")
#
#         # DB ga PDF saqlash
#         with open(pdf_path, "rb") as f:
#             contract_file = ContractFiles.objects.create(student=student)
#             contract_file.file.save(f"contract_{student.id}.pdf", f)
#
#         # Vaqtinchalik fayllarni o'chirish
#         os.remove(output_docx)
#         if os.path.exists(pdf_path):
#             os.remove(pdf_path)
#
#         return FileResponse(
#             contract_file.file.open("rb"),
#             as_attachment=True,
#             filename=f"contract_{student.id}.pdf"
#         )
#
#     def replace_text_preserving_format(self, paragraph, replacements):
#         """Formatlashni saqlab qolgan holda matnni almashtirish"""
#         # Paragraphdagi barcha runlarni tekshiramiz
#         for run in paragraph.runs:
#             original_text = run.text
#             new_text = original_text
#
#             # Har bir kalit so'zni almashtiramiz
#             for key, value in replacements.items():
#                 if key in new_text:
#                     new_text = new_text.replace(key, value)
#
#             # Agar matn o'zgarganga, runni yangilaymiz
#             if new_text != original_text:
#                 run.text = new_text
#
#
#
